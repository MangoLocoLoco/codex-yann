from io import BytesIO
from docx import Document
from docx.shared import Inches


def build_docx(pages, output_path):
    """Build a DOCX file from parsed PDF pages."""
    document = Document()
    for page_index, items in enumerate(pages):
        for item in items:
            if item.get("type") == "text":
                document.add_paragraph(item.get("text", ""))
            elif item.get("type") == "image":
                image_bytes = item.get("data")
                if image_bytes:
                    document.add_picture(BytesIO(image_bytes), width=Inches(6))
        if page_index != len(pages) - 1:
            document.add_page_break()
    document.save(output_path)
